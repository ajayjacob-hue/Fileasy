from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import os
import shutil
import uuid
import tempfile
from pathlib import Path

# Converters
from docx import Document
from pdf2docx import parse
from docx2pdf import convert as convert_word
from pypdf import PdfReader, PdfWriter
import fitz  # PyMuPDF

try:
    import comtypes.client
    COMTYPES_AVAILABLE = True
except ImportError:
    COMTYPES_AVAILABLE = False

app = FastAPI(title="Fileasy API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMP_DIR = Path(tempfile.gettempdir()) / "fileasy_temp"
TEMP_DIR.mkdir(exist_ok=True)

def cleanup_file(filepath: Path):
    if filepath.exists():
        try:
            os.remove(filepath)
        except Exception:
            pass

def fix_ligatures(docx_path):
    doc = Document(docx_path)
    replacements = {
        'Ɵ': 'ti',
        'Ʃ': 'tl',
        'Ō': 'ft',
        'ƞ': 'ti'
    }
    for p in doc.paragraphs:
        for run in p.runs:
            for bad, good in replacements.items():
                if bad in run.text:
                    run.text = run.text.replace(bad, good)
        if p.text.strip().startswith('?'):
            for run in p.runs:
                if '?' in run.text:
                    run.text = run.text.replace('?', '•', 1)
                    break
                    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for bad, good in replacements.items():
                            if bad in run.text:
                                run.text = run.text.replace(bad, good)
                    if p.text.strip().startswith('?'):
                        for run in p.runs:
                            if '?' in run.text:
                                run.text = run.text.replace('?', '•', 1)
                                break
    doc.save(docx_path)

@app.post("/api/pdf-to-word")
async def pdf_to_word(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files allowed")
    
    unique_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{unique_id}_input.pdf"
    output_path = TEMP_DIR / f"{unique_id}_output.docx"
    
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
        
    try:
        parse(str(input_path), str(output_path))
        fix_ligatures(str(output_path))
            
        background_tasks.add_task(cleanup_file, input_path)
        background_tasks.add_task(cleanup_file, output_path)
        return FileResponse(output_path, filename=f"{Path(file.filename).stem}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        cleanup_file(input_path)
        print("PDF2DOCX ERROR:", str(e))
        raise HTTPException(status_code=500, detail="File Error: " + str(e))

@app.post("/api/word-to-pdf")
async def word_to_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not (file.filename.endswith('.docx') or file.filename.endswith('.doc')):
        raise HTTPException(status_code=400, detail="Only Word files allowed")
    
    unique_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{unique_id}_{file.filename}"
    output_path = TEMP_DIR / f"{unique_id}_output.pdf"
    
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
        
    try:
        convert_word(str(input_path), str(output_path))
        background_tasks.add_task(cleanup_file, input_path)
        background_tasks.add_task(cleanup_file, output_path)
        return FileResponse(output_path, filename=f"{Path(file.filename).stem}.pdf", media_type="application/pdf")
    except Exception as e:
        cleanup_file(input_path)
        raise HTTPException(status_code=500, detail=f"Word to PDF requires MS Word installed on window. Error: {str(e)}")

@app.post("/api/ppt-to-pdf")
async def ppt_to_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not (file.filename.endswith('.pptx') or file.filename.endswith('.ppt')):
        raise HTTPException(status_code=400, detail="Only PPT files allowed")
        
    if not COMTYPES_AVAILABLE:
        raise HTTPException(status_code=500, detail="Comtypes is not installed or not available on this platform.")

    unique_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{unique_id}_{file.filename}"
    output_path = TEMP_DIR / f"{unique_id}_output.pdf"
    
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
        
    try:
        powerpoint = comtypes.client.CreateObject("Powerpoint.Application")
        powerpoint.Visible = 1
        
        # PPT to PDF format code is 32
        presentation = powerpoint.Presentations.Open(str(input_path), WithWindow=False)
        presentation.SaveAs(str(output_path), 32)
        presentation.Close()
        # Do not quit powerpoint, it might close user's other PPTs
        
        background_tasks.add_task(cleanup_file, input_path)
        background_tasks.add_task(cleanup_file, output_path)
        return FileResponse(output_path, filename=f"{Path(file.filename).stem}.pdf", media_type="application/pdf")
    except Exception as e:
        cleanup_file(input_path)
        raise HTTPException(status_code=500, detail=f"PPT conversion requires MS PowerPoint. Error: {str(e)}")

@app.post("/api/compress")
async def compress_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    unique_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{unique_id}_input.pdf"
    output_path = TEMP_DIR / f"{unique_id}_output.pdf"
    
    file_bytes = await file.read()
    with open(input_path, "wb") as f:
        f.write(file_bytes)
        
    try:
        # Deep PDF compression via PyMuPDF multi-pass image rewriting
        target_size = 4 * 1024 * 1024
        tmp_path = TEMP_DIR / f"{unique_id}_tmp.pdf"
        
        params = [
            {"dpi": 144, "quality": 50},
            {"dpi": 72, "quality": 25},
            {"dpi": 36, "quality": 10}
        ]
        
        current_input = input_path
        
        for i, param in enumerate(params):
            doc = fitz.open(current_input)
            if doc.is_encrypted:
                doc.authenticate("")
                
            try:
                if hasattr(doc, 'rewrite_images'):
                    doc.rewrite_images(dpi_target=param["dpi"], quality=param["quality"])
            except Exception as img_e:
                print(f"Image compression pass {i} skipped:", str(img_e))
                
            doc.save(output_path, garbage=4, deflate=True, clean=True)
            doc.close()
            
            if os.path.getsize(output_path) <= target_size:
                break
                
            # If not the last iteration, move output to tmp to read from it next pass
            if i < len(params) - 1:
                os.replace(output_path, tmp_path)
                current_input = tmp_path
                
        cleanup_file(tmp_path)
        
        background_tasks.add_task(cleanup_file, input_path)
        background_tasks.add_task(cleanup_file, output_path)
        return FileResponse(output_path, filename=f"{Path(file.filename).stem}_compressed.pdf", media_type="application/pdf")
    except Exception as e:
        cleanup_file(input_path)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/protect")
async def protect_pdf(background_tasks: BackgroundTasks, password: str = Form(...), file: UploadFile = File(...)):
    unique_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{unique_id}_input.pdf"
    output_path = TEMP_DIR / f"{unique_id}_output.pdf"
    
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
        
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
            
        writer.encrypt(password)
        with open(output_path, "wb") as f:
            writer.write(f)
            
        background_tasks.add_task(cleanup_file, input_path)
        background_tasks.add_task(cleanup_file, output_path)
        return FileResponse(output_path, filename=f"{Path(file.filename).stem}_protected.pdf", media_type="application/pdf")
    except Exception as e:
        cleanup_file(input_path)
        raise HTTPException(status_code=500, detail=str(e))

import pikepdf
from typing import Optional

@app.post("/api/unlock")
async def unlock_pdf(background_tasks: BackgroundTasks, password: Optional[str] = Form(None), file: UploadFile = File(...)):
    if not password:
        raise HTTPException(status_code=400, detail="Password is required to unlock this PDF.")
        
    unique_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{unique_id}_input.pdf"
    output_path = TEMP_DIR / f"{unique_id}_output.pdf"
    
    with open(input_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
        
    try:
        try:
            pdf = pikepdf.open(str(input_path), password=password)
            pdf.save(str(output_path))
            pdf.close()
        except pikepdf.PasswordError:
            raise Exception("Incorrect password provided.")
            
        background_tasks.add_task(cleanup_file, input_path)
        background_tasks.add_task(cleanup_file, output_path)
        return FileResponse(output_path, filename=f"{Path(file.filename).stem}_unlocked.pdf", media_type="application/pdf")
    except Exception as e:
        cleanup_file(input_path)
        raise HTTPException(status_code=500, detail=str(e))

import sys

# Dynamically resolve the absolute path to the static directory
# This handles the temporary _MEIPASS extraction folder when running as a PyInstaller .exe
if getattr(sys, 'frozen', False):
    base_dir = sys._MEIPASS
else:
    base_dir = os.path.dirname(os.path.abspath(__file__))

static_path = os.path.join(base_dir, "static")

# Mount the React Frontend build
app.mount("/", StaticFiles(directory=static_path, html=True), name="static")

@app.get("/api/shutdown")
def shutdown():
    import os
    import signal
    # Graceful exit
    os.kill(os.getpid(), signal.SIGTERM)
    return {"message": "Shutting down..."}

if __name__ == "__main__":
    import uvicorn
    import threading
    import webbrowser
    import time
    import urllib.request
    from multiprocessing import freeze_support

    # Fixes recursion bugs when packaging uvicorn with Pyinstaller
    freeze_support()

    def run_server():
        try:
            import sys
            import os
            # Fix for PyInstaller --noconsole removing standard output
            # Uvicorn crashes trying to colorize logs if sys.stdout is None
            if getattr(sys, 'stdout', None) is None:
                sys.stdout = open(os.devnull, "w")
            if getattr(sys, 'stderr', None) is None:
                sys.stderr = open(os.devnull, "w")
            if getattr(sys, 'stdin', None) is None:
                sys.stdin = open(os.devnull, "r")
                
            uvicorn.run(app, host="127.0.0.1", port=8000, log_config=None)
        except Exception as e:
            import traceback
            with open("Server_Crash_Log.txt", "w") as f:
                f.write(traceback.format_exc())

    t = threading.Thread(target=run_server, daemon=True)
    t.start()
    
    server_online = False
    print("Waiting for Fileasy Desktop Engine to secure a local port...")
    for _ in range(30):
        try:
            urllib.request.urlopen("http://127.0.0.1:8000")
            server_online = True
            break
        except Exception:
            time.sleep(0.5)

    if server_online:
        print("Launching Fileasy Desktop Environment...")
        webbrowser.open("http://127.0.0.1:8000")
        try:
            while True:
                time.sleep(1)
        except KeyboardInterrupt:
            pass
    else:
        # Write to crash log if even the thread didn't catch it
        with open("Server_Crash_Log.txt", "a") as f:
            f.write("\nServer failed to respond within 15 seconds.")
